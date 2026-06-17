from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


ROOT = Path("/Users/prathamgupta/Desktop/AttendanceSystem")
OUTPUT_DIR = ROOT / "artifacts" / "marketing"
LOGO_PATH = ROOT / "frontend" / "src" / "assets" / "peeplify-logo.png"
DOCX_PATH = OUTPUT_DIR / "Peeplify-Showcase-Brochure.docx"


PRIMARY = RGBColor(29, 78, 216)
PRIMARY_DARK = RGBColor(22, 48, 85)
TEXT_SOFT = RGBColor(88, 112, 143)
BG_SOFT = "EDF5FF"
BG_STRONG = "DBEAFE"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_page_borders(section):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "24")
        elem.set(qn("w:color"), "D6E5FF")
        pg_borders.append(elem)
    sect_pr.append(pg_borders)


def add_run(paragraph, text, *, bold=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = "Arial"
    return run


def add_feature_card(table, row, col, title, body):
    cell = table.cell(row, col)
    set_cell_shading(cell, BG_SOFT)
    set_cell_margins(cell, top=150, start=150, bottom=150, end=150)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, title + "\n", bold=True, size=12, color=PRIMARY_DARK)
    add_run(p, body, size=9.5, color=TEXT_SOFT)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    set_page_borders(section)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Inches(2.1)
    header.columns[1].width = Inches(5.2)

    logo_cell = header.cell(0, 0)
    set_cell_margins(logo_cell, top=80, start=80, bottom=80, end=80)
    if LOGO_PATH.exists():
        para = logo_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))

    hero_cell = header.cell(0, 1)
    set_cell_margins(hero_cell, top=110, start=140, bottom=110, end=120)
    hero = hero_cell.paragraphs[0]
    hero.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(hero, "PEEPLIFY\n", bold=True, size=24, color=PRIMARY)
    add_run(hero, "Workforce Management, Simplified\n", bold=True, size=15, color=PRIMARY_DARK)
    add_run(hero, "Attendance, payroll, leave, roster, proof and owner control in one easy mobile-first platform.", size=10.5, color=TEXT_SOFT)

    doc.add_paragraph()

    strap = doc.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(strap, "Built for Indian businesses with 10–50 employees", bold=True, size=11, color=PRIMARY_DARK)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        intro,
        "Peeplify helps owners reduce attendance disputes, track staff with selfie and GPS proof, manage leave and payroll, and run day-to-day workforce operations from phone or desktop.",
        size=10.5,
        color=TEXT_SOFT,
    )

    feature_heading = doc.add_paragraph()
    feature_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(feature_heading, "What Peeplify handles", bold=True, size=13, color=PRIMARY_DARK)

    feature_table = doc.add_table(rows=2, cols=3)
    feature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    feature_table.autofit = False
    for col in feature_table.columns:
        col.width = Inches(2.35)
    add_feature_card(feature_table, 0, 0, "Attendance with proof", "Selfie + location-backed check-in and check-out with employee profile image support.")
    add_feature_card(feature_table, 0, 1, "Leave and corrections", "Leave requests, missed punch fixes, owner approvals, and clean monthly records.")
    add_feature_card(feature_table, 0, 2, "Payroll clarity", "Salary snapshots, advance payments, payable days, and branch-specific workday rules.")
    add_feature_card(feature_table, 1, 0, "Roster management", "Shift setup, templates, weekly offs, planning, conflicts, and assignment control.")
    add_feature_card(feature_table, 1, 1, "Owner-first dashboard", "See who is present, absent, late, or forgot checkout at a glance.")
    add_feature_card(feature_table, 1, 2, "Commercial-ready setup", "Plans, trial expiry, billing, PDF invoices, support requests, and renewal flow.")

    industries_heading = doc.add_paragraph()
    industries_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(industries_heading, "Ideal for these businesses", bold=True, size=13, color=PRIMARY_DARK)

    industries = doc.add_table(rows=2, cols=2)
    industries.alignment = WD_TABLE_ALIGNMENT.CENTER
    industries.autofit = False
    industries.columns[0].width = Inches(3.45)
    industries.columns[1].width = Inches(3.45)
    left_items = [
        "Hospitals, nursing homes, clinics, diagnostic centres",
        "Hotels, guest houses, restaurants, banquet halls",
        "Retail shops, supermarkets, showrooms",
        "Pharmacies, salons, gyms, fitness centres",
    ]
    right_items = [
        "Factories, workshops, petrol pumps, industrial plants",
        "Security agencies, housekeeping services",
        "Schools, coaching institutes, training centres",
        "IT companies, consultancies and offices",
    ]
    for idx, items in enumerate((left_items, right_items)):
        cell = industries.cell(0, idx)
        set_cell_shading(cell, BG_SOFT)
        set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
        for i, item in enumerate(items):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            add_run(p, f"• {item}", size=10, color=PRIMARY_DARK)
    set_cell_shading(industries.cell(1, 0), BG_STRONG)
    set_cell_margins(industries.cell(1, 0), top=130, start=160, bottom=130, end=160)
    p = industries.cell(1, 0).paragraphs[0]
    add_run(p, "Target market\n", bold=True, size=11.5, color=PRIMARY_DARK)
    add_run(p, "Punjab and North India SMBs that want stronger discipline without enterprise complexity.", size=9.5, color=TEXT_SOFT)
    set_cell_shading(industries.cell(1, 1), BG_STRONG)
    set_cell_margins(industries.cell(1, 1), top=130, start=160, bottom=130, end=160)
    p = industries.cell(1, 1).paragraphs[0]
    add_run(p, "Why owners choose it\n", bold=True, size=11.5, color=PRIMARY_DARK)
    add_run(p, "Simple setup, mobile-first workflow, clearer payroll, and fewer attendance disputes.", size=9.5, color=TEXT_SOFT)

    pricing_heading = doc.add_paragraph()
    pricing_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(pricing_heading, "Simple pricing", bold=True, size=13, color=PRIMARY_DARK)

    pricing = doc.add_table(rows=2, cols=3)
    pricing.alignment = WD_TABLE_ALIGNMENT.CENTER
    pricing.autofit = False
    for col in pricing.columns:
        col.width = Inches(2.28)
    plans = [
        ("Starter", "Up to 10 employees", "INR 1199 / month"),
        ("Growth", "Up to 25 employees", "INR 1999 / month"),
        ("Business", "Up to 50 employees + multiple branches", "INR 2999 / month"),
    ]
    for idx, (title, sub, price) in enumerate(plans):
        cell = pricing.cell(0, idx)
        set_cell_shading(cell, BG_SOFT if idx != 2 else BG_STRONG)
        set_cell_margins(cell, top=140, start=150, bottom=150, end=150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, title + "\n", bold=True, size=12.5, color=PRIMARY_DARK)
        add_run(p, sub + "\n", size=9.5, color=TEXT_SOFT)
        add_run(p, price, bold=True, size=14, color=PRIMARY)
    cell = pricing.cell(1, 0)
    cell.merge(pricing.cell(1, 2))
    set_cell_shading(cell, BG_SOFT)
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Quarterly, half-yearly, and yearly billing discounts available. 3-day free access for new workspaces.", size=10, color=PRIMARY_DARK)

    cta = doc.add_paragraph()
    cta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cta, "Book a demo or start your workspace today", bold=True, size=13, color=PRIMARY_DARK)

    footer = doc.add_table(rows=1, cols=3)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer.autofit = False
    footer.columns[0].width = Inches(2.2)
    footer.columns[1].width = Inches(2.9)
    footer.columns[2].width = Inches(1.8)
    footer_data = [
        ("Website", "www.peeplify.com"),
        ("Email", "pratham4714@gmail.com"),
        ("Pitch", "Track people, time and payouts in one place"),
    ]
    for idx, (label, value) in enumerate(footer_data):
        cell = footer.cell(0, idx)
        set_cell_shading(cell, BG_SOFT)
        set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, label + "\n", bold=True, size=10.5, color=PRIMARY_DARK)
        add_run(p, value, size=9.3, color=TEXT_SOFT)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
